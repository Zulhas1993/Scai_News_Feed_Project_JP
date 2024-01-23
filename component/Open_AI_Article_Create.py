import json
from azure.ai.textanalytics import TextAnalyticsClient
from azure.core.credentials import AzureKeyCredential
from Details_News import generate_news_details_list
from bs4 import BeautifulSoup
from docx import Document

def read_json(file_path):
    with open(file_path, 'r', encoding='utf-16') as file:
        data = json.load(file)
    return data

def authenticate_client():
    endpoint = "https://laboblogtextanalytics.cognitiveservices.azure.com/"
    key = "09b7c88cfff44bac95c83a2256f31907"
    return TextAnalyticsClient(endpoint=endpoint, credential=AzureKeyCredential(key))

def generate_abstractive_summary(text_analytics_client, document_to_summarize):
    try:
        # Begin abstractive summarization
        poller = text_analytics_client.begin_abstractive_summarization(documents=[{"id": "1", "text": document_to_summarize}])
        abstract_summary_results = poller.result()

        summaries = []
        for result in abstract_summary_results:
            if result.kind == "AbstractiveSummarization":
                summaries.append(result.summaries)

        return summaries

    except Exception as e:
        print(f"An error occurred during abstractive summarization: {e}")
        return None

def generate_and_save_summary_docx(text_analytics_client):
    try:
        # Generate news details list
        news_details_list = generate_news_details_list()

        # Create a DOCX document
        doc = Document()

        for news_details in news_details_list:
            title = news_details.get('title', '')
            link = news_details.get('link', '')
            details = news_details.get('details', {}).get('content', '')

            # Generate abstractive summary
            summaries = generate_abstractive_summary(text_analytics_client, details)

            # Add news details and summaries to the document
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"Link: {link}")
            doc.add_paragraph("Details:")
            doc.add_paragraph(details)

            if summaries:
                doc.add_paragraph("Abstractive Summary:")
                for summary in summaries:
                    doc.add_paragraph(summary.text)

            doc.add_page_break()

        # Save the document
        doc.save('news_summaries.docx')
        print("News summaries saved to news_summaries.docx")

    except Exception as e:
        print(f"An error occurred while generating and saving the summary DOCX: {e}")

# Authenticate the Text Analytics client
text_analytics_client = authenticate_client()

# Generate and save abstractive summaries in DOCX
generate_and_save_summary_docx(text_analytics_client)
